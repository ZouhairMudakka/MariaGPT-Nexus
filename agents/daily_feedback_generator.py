import os
from datetime import datetime, timedelta
from pathlib import Path
import json
from typing import Dict, List, Any
from docx import Document
import logging
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import smtplib

logger = logging.getLogger(__name__)

class DailyFeedbackGenerator:
    def __init__(self, evaluations_path: str = "data/agent_evaluations"):
        self.base_path = Path(evaluations_path)
        self.metrics_path = self.base_path / "metrics"
        self.feedback_path = self.base_path / "daily_feedback"
        self.feedback_path.mkdir(parents=True, exist_ok=True)

    def generate_daily_feedback(self) -> None:
        """Generate daily feedback document based on last 24h evaluations"""
        try:
            # Get last 24h evaluations
            performance_data = self._load_recent_evaluations(
                self.metrics_path / "agent_performance_metrics.json")
            flow_data = self._load_recent_evaluations(
                self.metrics_path / "conversation_flow_metrics.json")
            
            if not performance_data or not flow_data:
                logger.warning("No evaluation data available for the last 24 hours")
                return
            
            # Create feedback document
            doc = Document()
            doc.add_heading('Daily Agent Performance Feedback', 0)
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            
            try:
                self._add_performance_insights(doc, performance_data)
                self._add_flow_insights(doc, flow_data)
                self._add_improvement_recommendations(doc, performance_data, flow_data)
            except Exception as e:
                logger.error(f"Error generating document content: {str(e)}")
                raise
            
            # Save document
            filename = f'agent_feedback_{datetime.now().strftime("%Y%m%d")}.docx'
            filepath = self.feedback_path / filename
            doc.save(filepath)
            
            # Send email with attachment
            self._send_feedback_email(filepath)
            
        except Exception as e:
            logger.error(f"Error generating daily feedback: {str(e)}")
            raise

    def _load_recent_evaluations(self, file_path: Path) -> Dict[str, Any]:
        """Load evaluations from last 24 hours"""
        try:
            with open(file_path, 'r') as f:
                data = json.load(f)
            
            cutoff_time = (datetime.now() - timedelta(days=1)).isoformat()
            
            if 'conversations' in data:
                data['conversations']['completed'] = [
                    conv for conv in data['conversations']['completed']
                    if conv['metadata']['timestamp'] > cutoff_time
                ]
            
            return data
        except Exception as e:
            logger.error(f"Error loading recent evaluations: {str(e)}")
            return {}

    def _add_performance_insights(self, doc: Document, performance_data: Dict[str, Any]) -> None:
        """Add performance insights section to document"""
        doc.add_heading('Performance Insights', level=1)
        
        completed_convs = performance_data.get('conversations', {}).get('completed', [])
        if not completed_convs:
            doc.add_paragraph('No conversations completed in the last 24 hours.')
            return
            
        # Analyze metrics
        avg_metrics = self._calculate_average_metrics(completed_convs)
        
        doc.add_paragraph('Average Performance Metrics:')
        for metric, value in avg_metrics.items():
            doc.add_paragraph(f'- {metric.replace("_", " ").title()}: {value:.2f}/10')

    def _add_flow_insights(self, doc: Document, flow_data: Dict[str, Any]) -> None:
        """Add conversation flow insights section"""
        doc.add_heading('Conversation Flow Analysis', level=1)
        
        flow_metrics = flow_data.get('flow_metrics', [])
        recent_flows = [
            flow for flow in flow_metrics
            if flow['timestamp'] > (datetime.now() - timedelta(days=1)).isoformat()
        ]
        
        if not recent_flows:
            doc.add_paragraph('No flow data available for the last 24 hours.')
            return
            
        # Analyze flow patterns
        doc.add_paragraph('Key Observations:')
        self._analyze_flow_patterns(doc, recent_flows)

    def _add_improvement_recommendations(self, doc: Document, 
                                      performance_data: Dict[str, Any],
                                      flow_data: Dict[str, Any]) -> None:
        """Add recommendations based on analysis"""
        doc.add_heading('Recommendations for Improvement', level=1)
        
        # Generate recommendations
        recommendations = self._generate_recommendations(performance_data, flow_data)
        
        for category, items in recommendations.items():
            doc.add_heading(category, level=2)
            for item in items:
                doc.add_paragraph(f'• {item}')

    def _send_feedback_email(self, filepath: Path) -> None:
        """Send email with feedback document attached"""
        try:
            msg = MIMEMultipart()
            msg['Subject'] = f'Daily Agent Performance Feedback - {datetime.now().strftime("%Y-%m-%d")}'
            msg['From'] = "maria.gpt@mudakka.com"
            msg['To'] = "hr@mudakka.com"
            
            # Add body
            body = """
            Daily Agent Performance Feedback Report
            
            This automated report contains:
            - Performance metrics analysis
            - Conversation flow insights
            - Improvement recommendations
            - Team collaboration metrics
            
            Please review the attached document for detailed analysis.
            """
            msg.attach(MIMEText(body, 'plain'))
            
            # Add attachment
            with open(filepath, 'rb') as f:
                attachment = MIMEApplication(f.read(), _subtype='docx')
                attachment.add_header('Content-Disposition', 'attachment', 
                                   filename=filepath.name)
                msg.attach(attachment)
            
            # Configure SMTP settings
            smtp_server = os.environ.get("SMTP_SERVER", "smtp.gmail.com")
            smtp_port = int(os.environ.get("SMTP_PORT", "587"))
            smtp_username = os.environ.get("SMTP_USERNAME", "maria.gpt@mudakka.com")
            smtp_password = os.environ.get("SMTP_PASSWORD")
            
            if not smtp_password:
                raise ValueError("SMTP password not configured")
            
            # Send email
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()
                server.login(smtp_username, smtp_password)
                server.send_message(msg)
                logger.info(f"Feedback email sent successfully to hr@mudakka.com")
                
        except Exception as e:
            logger.error(f"Error sending feedback email: {str(e)}")
            raise

    def _calculate_average_metrics(self, conversations: List[Dict[str, Any]]) -> Dict[str, float]:
        """Calculate average metrics from conversations"""
        metrics_sum = {}
        metrics_count = {}
        
        for conv in conversations:
            for metric, value in conv["metrics"]["conversation"].items():
                if isinstance(value, (int, float)):
                    metrics_sum[metric] = metrics_sum.get(metric, 0) + value
                    metrics_count[metric] = metrics_count.get(metric, 0) + 1
        
        return {
            metric: metrics_sum[metric] / count
            for metric, count in metrics_count.items()
        }

    def _analyze_flow_patterns(self, doc: Document, flow_metrics: List[Dict[str, Any]]) -> None:
        """Analyze and document conversation flow patterns"""
        avg_handoff = sum(f["flow_metrics"]["handoff_smoothness"] for f in flow_metrics) / len(flow_metrics)
        avg_coherence = sum(f["flow_metrics"]["coherence"] for f in flow_metrics) / len(flow_metrics)
        avg_context = sum(f["flow_metrics"]["context_retention"] for f in flow_metrics) / len(flow_metrics)
        
        doc.add_paragraph(f'• Average Handoff Smoothness: {avg_handoff:.2f}/10')
        doc.add_paragraph(f'• Conversation Coherence: {avg_coherence:.2f}/10')
        doc.add_paragraph(f'• Context Retention: {avg_context:.2f}/10')
        
        # Analyze interaction patterns
        total_switches = sum(f["interaction_pattern"]["agent_switches"] for f in flow_metrics)
        total_transitions = sum(f["interaction_pattern"]["topic_transitions"] for f in flow_metrics)
        
        doc.add_paragraph(f'• Total Agent Switches: {total_switches}')
        doc.add_paragraph(f'• Total Topic Transitions: {total_transitions}')

    def _generate_recommendations(self, performance_data: Dict[str, Any], 
                            flow_data: Dict[str, Any]) -> Dict[str, List[str]]:
        """Generate improvement recommendations based on metrics analysis"""
        recommendations = {
            "Agent Performance": [],
            "Conversation Flow": [],
            "User Experience": [],
            "Team Collaboration": []
        }
        
        completed_convs = performance_data.get('conversations', {}).get('completed', [])
        if completed_convs:
            for conv in completed_convs:
                metrics = conv["metrics"]["conversation"]
                
                # Check flow metrics
                if metrics.get("flow_quality", 0) < 7:
                    recommendations["Conversation Flow"].append(
                        "Improve conversation flow and natural progression"
                    )
                if metrics.get("context_consistency", 0) < 7:
                    recommendations["Conversation Flow"].append(
                        "Enhance context maintenance throughout conversations"
                    )
                    
                # Check efficiency metrics
                if metrics.get("time_efficiency", 0) < 7:
                    recommendations["Agent Performance"].append(
                        "Work on reducing resolution time"
                    )
                if metrics.get("resource_utilization", 0) < 7:
                    recommendations["Team Collaboration"].append(
                        "Optimize agent resource allocation"
                    )
                    
                # Check outcome metrics
                if metrics.get("user_satisfaction", 0) < 7:
                    recommendations["User Experience"].append(
                        "Focus on improving overall user satisfaction"
                    )
                if metrics.get("goal_achievement", 0) < 7:
                    recommendations["User Experience"].append(
                        "Enhance goal achievement rate"
                    )